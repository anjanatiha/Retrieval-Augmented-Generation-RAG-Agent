"""test_file_upload_tools.py — All 5 agent tools exercised with real uploaded file content.

Tests that every agent tool (rag_search, calculator, summarise, sentiment, finish)
works correctly when the underlying content comes from real temp files processed by
DocumentLoader._dispatch_chunker.

Mock strategy:
  Always mock: ollama.embed, ollama.chat, chromadb → EphemeralClient
  Never mock:  fitz, python-docx, openpyxl, python-pptx, BM25Okapi, file I/O,
               calculator eval

Reason for split: max 500 lines per file per CLAUDE.md.
Chat mode and agent mode file upload tests are in test_file_upload.py.
"""

import uuid
import os
import pytest
import chromadb
from rank_bm25 import BM25Okapi
from unittest.mock import patch, MagicMock


# ---------------------------------------------------------------------------
# Shared helpers (duplicated from test_file_upload.py — each file is standalone)
# ---------------------------------------------------------------------------

def _fake_embed(**kwargs):
    """Return a fixed 4-dim embedding vector to satisfy ollama.embed calls."""
    return {'embeddings': [[0.1, 0.2, 0.3, 0.4]]}


def _fake_chat(**kwargs):
    """Return a canned chat response (stream or single) to satisfy ollama.chat calls."""
    if kwargs.get('stream'):
        return [{'message': {'content': 'mock response'}}]
    return {'message': {'content': 'mock response'}}


def _make_store(chunks):
    """Build an isolated VectorStore with EphemeralClient."""
    from src.rag.vector_store import VectorStore
    vs         = VectorStore()
    client     = chromadb.EphemeralClient()
    collection = client.get_or_create_collection(
        f'test_{uuid.uuid4().hex}', metadata={'hnsw:space': 'cosine'}
    )
    if chunks:
        ids    = [f'c{i}' for i in range(len(chunks))]
        texts  = [c['text'] for c in chunks]
        metas  = [{'source': c['source'], 'start_line': c['start_line'],
                   'end_line': c['end_line'], 'type': c['type']} for c in chunks]
        embeds = [[0.1, 0.2, 0.3, 0.4]] * len(chunks)
        collection.add(ids=ids, embeddings=embeds, documents=texts, metadatas=metas)
    vs.collection           = collection
    vs.chunks               = chunks
    vs.bm25_index           = BM25Okapi([c['text'].lower().split() for c in chunks]) if chunks else None
    vs.conversation_history = []
    return vs


PATCHES = [
    patch('ollama.embed', side_effect=_fake_embed),
    patch('ollama.chat',  side_effect=_fake_chat),
]


def _apply_patches(func):
    """Decorator that applies both ollama.embed and ollama.chat patches to a test method."""
    for p in reversed(PATCHES):
        func = p(func)
    return func


def _upload_chunks(tmp_path, filename, content_fn):
    """Create a real temp file and run it through DocumentLoader._dispatch_chunker."""
    from src.rag.document_loader import DocumentLoader
    loader   = DocumentLoader()
    ext      = os.path.splitext(filename)[1].lower()
    dtype    = loader.ext_to_type.get(ext, 'txt')
    filepath = str(tmp_path / filename)
    content_fn(filepath)
    file_info = {
        'filepath':      filepath,
        'filename':      filename,
        'detected_type': dtype,
        'is_misplaced':  False,
    }
    return loader._dispatch_chunker(file_info)


# ---------------------------------------------------------------------------
# Per-type file builders — write real files using the actual libraries
# ---------------------------------------------------------------------------

def _write_pdf(path):
    """Single-page PDF with nutrition facts."""
    import fitz
    doc  = fitz.open(); page = doc.new_page()
    page.insert_text((50, 50),
        'Nutrition facts: Bananas contain 89 calories per 100 grams. '
        'They are rich in potassium and vitamin B6. '
        'Avocados contain 160 calories per 100 grams and healthy fats.')
    doc.save(path); doc.close()


def _write_txt(path):
    """Plain-text dog facts, one fact per line."""
    with open(path, 'w') as f:
        f.write('Dogs were domesticated from wolves approximately 15000 years ago.\n')
        f.write('There are over 340 recognized dog breeds in the world.\n')
        f.write('A dog can understand about 165 words and gestures.\n')


def _write_docx(path):
    """DOCX with three programming language paragraphs."""
    from docx import Document
    doc = Document()
    doc.add_paragraph('Python is a programming language created by Guido van Rossum in 1991.')
    doc.add_paragraph('Java was created by James Gosling at Sun Microsystems in 1995.')
    doc.add_paragraph('Rust was developed at Mozilla Research and released in 2010.')
    doc.save(path)


def _write_xlsx(path):
    """XLSX with country/capital/population data."""
    import openpyxl
    wb = openpyxl.Workbook(); ws = wb.active
    ws.append(['Country', 'Capital', 'Population_Millions'])
    ws.append(['Brazil',  'Brasilia', 215])
    ws.append(['Japan',   'Tokyo',    125])
    ws.append(['Germany', 'Berlin',   84])
    wb.save(path)


def _write_csv(path):
    """CSV with programming language creator and year."""
    with open(path, 'w') as f:
        f.write('language,year,creator\n')
        f.write('Python,1991,Guido van Rossum\n')
        f.write('Rust,2010,Graydon Hoare\n')
        f.write('Go,2009,Robert Griesemer\n')


def _write_md(path):
    """Markdown file with coffee facts."""
    with open(path, 'w') as f:
        f.write('# Coffee Facts\n\n')
        f.write("Coffee was first discovered in Ethiopia around 850 AD by a goat herder named Kaldi.\n\n")
        f.write("Brazil is the world's largest coffee producer, accounting for 40 percent of global supply.\n\n")
        f.write('A standard cup of coffee contains between 80 and 100 milligrams of caffeine.\n')


def _write_html(path):
    """HTML file with dog biology facts."""
    with open(path, 'w') as f:
        f.write('<html><body><h1>Dog Facts</h1>')
        f.write("<p>A dog's sense of smell is up to 100000 times more powerful than a human's.</p>")
        f.write('<p>Border Collies are considered the most intelligent dog breed.</p>')
        f.write('<p>Greyhounds can run at speeds up to 72 kilometers per hour.</p>')
        f.write('</body></html>')


# ---------------------------------------------------------------------------
# 3. All 5 agent tools with real uploaded file content
# ---------------------------------------------------------------------------

class TestFileUploadAllTools:
    """All 5 tools (rag_search, calculator, summarise, sentiment, finish) exercised
    with content loaded from real temp files via DocumentLoader._dispatch_chunker."""

    @_apply_patches
    def test_rag_search_on_txt(self, mock_chat, mock_embed, tmp_path):
        """rag_search returns chunks from a real uploaded txt file."""
        chunks = _upload_chunks(tmp_path, 'dogs.txt', _write_txt)
        from src.rag.agent import Agent
        result = Agent(_make_store(chunks))._tool_rag_search('dog breeds')
        assert len(result) > 0 and 'dogs.txt' in result

    @_apply_patches
    def test_rag_search_on_csv(self, mock_chat, mock_embed, tmp_path):
        """rag_search returns chunks from a real uploaded csv file."""
        chunks = _upload_chunks(tmp_path, 'languages.csv', _write_csv)
        from src.rag.agent import Agent
        result = Agent(_make_store(chunks))._tool_rag_search('Python creator')
        assert len(result) > 0

    @_apply_patches
    def test_calculator_on_xlsx_values(self, mock_chat, mock_embed, tmp_path):
        """calculator evaluates numbers that might come from an xlsx file."""
        chunks = _upload_chunks(tmp_path, 'countries.xlsx', _write_xlsx)
        from src.rag.agent import Agent
        assert Agent(_make_store(chunks))._tool_calculator('215 + 125 + 84') == '424'

    @_apply_patches
    def test_calculator_percentage(self, mock_chat, mock_embed, tmp_path):
        """calculator handles percentage expressions."""
        from src.rag.agent import Agent
        result = Agent(_make_store([]))._tool_calculator('15% of 89')
        assert float(result) == pytest.approx(13.35)

    @_apply_patches
    def test_summarise_on_md(self, mock_chat, mock_embed, tmp_path):
        """summarise tool works on text from a real md file."""
        mock_chat.side_effect = lambda **kw: {'message': {'content': 'Coffee originated in Ethiopia.'}}
        chunks = _upload_chunks(tmp_path, 'coffee.md', _write_md)
        text   = ' '.join(c['text'] for c in chunks)
        from src.rag.agent import Agent
        result = Agent(_make_store(chunks))._tool_summarise(text)
        assert isinstance(result, str) and len(result) > 0

    @_apply_patches
    def test_summarise_on_docx(self, mock_chat, mock_embed, tmp_path):
        """summarise tool works on text from a real docx file."""
        mock_chat.side_effect = lambda **kw: {'message': {'content': 'Python was created by Guido.'}}
        chunks = _upload_chunks(tmp_path, 'languages.docx', _write_docx)
        text   = ' '.join(c['text'] for c in chunks)
        from src.rag.agent import Agent
        result = Agent(_make_store(chunks))._tool_summarise(text)
        assert isinstance(result, str) and len(result) > 0

    @_apply_patches
    def test_sentiment_on_html(self, mock_chat, mock_embed, tmp_path):
        """sentiment tool returns all 4 required fields for html file content."""
        resp = ('Sentiment: Positive\nTone: informative\n'
                'Key phrases: powerful, intelligent\nExplanation: Factual and positive.')
        mock_chat.side_effect = lambda **kw: {'message': {'content': resp}}
        chunks = _upload_chunks(tmp_path, 'dogs.html', _write_html)
        text   = ' '.join(c['text'] for c in chunks)
        from src.rag.agent import Agent
        result = Agent(_make_store(chunks))._tool_sentiment(text)
        assert 'Sentiment' in result and 'Tone' in result
        assert 'Key phrases' in result and 'Explanation' in result

    @_apply_patches
    def test_sentiment_on_txt(self, mock_chat, mock_embed, tmp_path):
        """sentiment tool works on text from a real txt file."""
        resp = ('Sentiment: Neutral\nTone: factual\n'
                'Key phrases: domesticated, breeds\nExplanation: Objective facts.')
        mock_chat.side_effect = lambda **kw: {'message': {'content': resp}}
        chunks = _upload_chunks(tmp_path, 'dogs.txt', _write_txt)
        text   = ' '.join(c['text'] for c in chunks)
        from src.rag.agent import Agent
        result = Agent(_make_store(chunks))._tool_sentiment(text)
        assert 'Sentiment' in result

    @_apply_patches
    def test_finish_after_rag_search(self, mock_chat, mock_embed, tmp_path):
        """finish step is reached after rag_search on a real uploaded file."""
        responses = iter([
            {'message': {'content': 'TOOL: rag_search(coffee Ethiopia)'}},
            {'message': {'content': 'TOOL: finish(Coffee was discovered in Ethiopia around 850 AD.)'}},
        ])
        mock_chat.side_effect = lambda **kw: next(responses)
        chunks = _upload_chunks(tmp_path, 'coffee.md', _write_md)
        from src.rag.agent import Agent
        result = Agent(_make_store(chunks)).run('Where was coffee discovered?', streamlit_mode=True)
        assert any(s['tool'] == 'finish' for s in result['steps'])
        assert result['answer'] is not None

    @_apply_patches
    def test_finish_calc_auto(self, mock_chat, mock_embed, tmp_path):
        """calculator auto-finish: tool result goes directly to finish."""
        mock_chat.side_effect = lambda **kw: {'message': {'content': 'TOOL: calculator(215 + 125)'}}
        from src.rag.agent import Agent
        result = Agent(_make_store([])).run('What is 215 plus 125?', streamlit_mode=True)
        assert any(s['tool'] == 'calculator' for s in result['steps'])
        assert any(s['tool'] == 'finish' for s in result['steps'])
        assert '340' in result['answer']

    @_apply_patches
    def test_all_5_tools_smoke(self, mock_chat, mock_embed, tmp_path):
        """Smoke test: all 5 tools are callable in sequence with real file content."""
        from src.rag.agent import Agent

        chunks = _upload_chunks(tmp_path, 'dogs.txt', _write_txt)
        store  = _make_store(chunks)
        agent  = Agent(store)

        # 1. rag_search
        assert len(agent._tool_rag_search('dog breeds')) > 0

        # 2. calculator
        assert agent._tool_calculator('340 * 2') == '680'

        # 3. summarise
        mock_chat.side_effect = lambda **kw: {'message': {'content': 'Dogs were domesticated long ago.'}}
        assert len(agent._tool_summarise(' '.join(c['text'] for c in chunks))) > 0

        # 4. sentiment
        mock_chat.side_effect = lambda **kw: {
            'message': {'content': 'Sentiment: Neutral\nTone: factual\nKey phrases: dogs\nExplanation: OK.'}
        }
        assert 'Sentiment' in agent._tool_sentiment(' '.join(c['text'] for c in chunks))

        # 5. finish — via auto-finish after rag_search
        mock_chat.side_effect = lambda **kw: {'message': {'content': 'TOOL: rag_search(dog breeds)'}}
        result = Agent(store).run('How many dog breeds are there?', streamlit_mode=True)
        assert any(s['tool'] == 'finish' for s in result['steps'])
